import os, cv2, json, numpy as np, pandas as pd
from ultralytics import YOLO
from tqdm import tqdm
from docx import Document

# =========================
# Settings
# =========================
VIDEO_PATH    = r"D:\research\clip_1.mp4"
MODEL_PATH    = r"D:\research\MOT20\weights\best.pt"
OUTPUT_DIR    = "id_mappings_MOT20"

OUTPUT_JSON   = os.path.join(OUTPUT_DIR, "frame_ids.json")
OUTPUT_DOCX   = os.path.join(OUTPUT_DIR, "frame_ids_report.docx")
OUTPUT_CSV    = os.path.join(OUTPUT_DIR, "ground_truth.csv")

CONF_THRESHOLD       = 0.5
SAVE_INTERVAL        = 5        # auto-save every N frames
DIST_THRESH_PX       = 50       # carry-forward match threshold (px)
HIGHLIGHT_BLINKS     = 2
SAVE_LABELED_FRAMES  = False    # True করলে লেবেল করা ইমেজও সেভ হবে

# =========================
# Init
# =========================
os.makedirs(OUTPUT_DIR, exist_ok=True)
if SAVE_LABELED_FRAMES:
    os.makedirs(os.path.join(OUTPUT_DIR, "labeled_frames"), exist_ok=True)

model = YOLO(MODEL_PATH)
model.classes = [0]  # person only

frame_id_data = {}    # {"12":[{"id":"A","x":..,"y":..,"box":[x1,y1,x2,y2]}, ...], ...}
id_colors     = {}    # persistent colors per ID
gt_rows       = []    # ground truth rows for CSV

def color_for_id(pid):
    if pid not in id_colors:
        np.random.seed(abs(hash(str(pid))) % (2**32))
        id_colors[pid] = tuple(np.random.randint(60, 255, 3).tolist())
    return id_colors[pid]

def random_color(seed=None):
    if seed is not None:
        np.random.seed(seed)
    return tuple(np.random.randint(60, 255, 3).tolist())

def detect_persons(frame):
    r = model(frame, verbose=False)
    boxes = r[0].boxes.xyxy.cpu().numpy()
    confs = r[0].boxes.conf.cpu().numpy() if hasattr(r[0].boxes, "conf") else np.ones(len(boxes))
    dets = []
    for i, b in enumerate(boxes):
        if confs[i] < CONF_THRESHOLD:
            continue
        x1, y1, x2, y2 = map(int, b[:4])
        if x2 <= x1 or y2 <= y1: 
            continue
        cx, cy = int((x1+x2)/2), int((y1+y2)/2)
        dets.append({"cx":cx, "cy":cy, "box":[x1,y1,x2,y2]})
    return dets

def nearest_match(prev_map, curr_pts, dist_thresh):
    """
    prev_map: {pid:(x,y)}
    curr_pts: [{"cx","cy","box"}, ...]
    Return: [{"cx","cy","box","pid" or None}, ...]
    """
    used = set()
    result = []
    for det in curr_pts:
        cx, cy = det["cx"], det["cy"]
        best_pid, best_d = None, 1e12
        for pid,(px,py) in prev_map.items():
            if pid in used: 
                continue
            d = np.hypot(cx - px, cy - py)
            if d < best_d and d <= dist_thresh:
                best_d, best_pid = d, pid
        if best_pid is not None:
            used.add(best_pid)
        out = det.copy()
        out["pid"] = best_pid
        result.append(out)
    return result

# =========================
# Video Loop
# =========================
cap = cv2.VideoCapture(VIDEO_PATH)
total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
frame_idx = 0

print("🟢 Smart Semi-Auto Labeling + Ground Truth (v10)")
print("Keys: [Enter]/[N]=Accept & label unmatched | [B]=Back | [Q]=Quit\n")

pbar = tqdm(total=total_frames, desc="Progress", ncols=80)

while True:
    if frame_idx < 0: frame_idx = 0
    cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
    ret, frame = cap.read()
    if not ret:
        print("📹 End of video.")
        break

    detections = detect_persons(frame)
    print(f"\n🟨 Frame {frame_idx}/{total_frames} | Found {len(detections)} persons")
    for i,d in enumerate(detections):
        print(f"  Person {i+1}: Center=({d['cx']},{d['cy']})")

    # Header overlay
    header = "Frame {}/{} | [Enter]/[N]=Next  [B]=Back  [Q]=Quit  | Tip: Enter w/o ID → Skip".format(frame_idx, total_frames)
    cv2.rectangle(frame, (10,10), (10+1000, 10+40), (0,0,0), -1)
    cv2.putText(frame, header, (20,38), cv2.FONT_HERSHEY_SIMPLEX, 0.7, (0,255,255), 2)

    # Prev frame map (for carry-forward)
    prev_map = {}
    if str(frame_idx-1) in frame_id_data:
        for o in frame_id_data[str(frame_idx-1)]:
            prev_map[str(o["id"])] = (int(o["x"]), int(o["y"]))

    assigned = nearest_match(prev_map, detections, DIST_THRESH_PX) if prev_map else [
        {**d, "pid": None} for d in detections
    ]

    # Draw: auto-IDs in their color; others as UNLABELED temp color
    temp_colors = [random_color(seed=i+frame_idx) for i in range(len(assigned))]
    for i, det in enumerate(assigned):
        x1,y1,x2,y2 = det["box"]
        if det["pid"]:
            col = color_for_id(det["pid"])
            cv2.rectangle(frame, (x1,y1),(x2,y2), col, 2)
            cv2.putText(frame, f"ID {det['pid']} (auto)", (x1+5,y1-8), cv2.FONT_HERSHEY_SIMPLEX, 0.7, col, 2)
        else:
            col = temp_colors[i]
            cv2.rectangle(frame, (x1,y1),(x2,y2), col, 2)
            cv2.putText(frame, "UNLABELED", (x1+5,y1-8), cv2.FONT_HERSHEY_SIMPLEX, 0.6, col, 2)

    cv2.imshow("Semi-Auto Labeling (v10)", frame)
    key = cv2.waitKey(0)

    if key == ord('q'):
        print("🟥 Quit."); break
    if key == ord('b'):
        frame_idx -= 1
        continue

    if key in (13, ord('n')):  # Enter or 'n'
        manual_entries = []

        # 1) keep auto-assigned (carry-forward) IDs
        for det in assigned:
            if det["pid"]:
                pid = str(det["pid"])
                x1,y1,x2,y2 = det["box"]
                col = color_for_id(pid)
                cv2.rectangle(frame,(x1,y1),(x2,y2),col,2)
                cv2.putText(frame,f"ID {pid}",(x1+5,y1-8),cv2.FONT_HERSHEY_SIMPLEX,0.7,col,2)
                manual_entries.append({"id": pid, "x": det["cx"], "y": det["cy"], "box": [x1,y1,x2,y2]})

        # 2) label new/unmatched ones (skip allowed)
        for det in assigned:
            if det["pid"] is None:
                x1,y1,x2,y2 = det["box"]

                # highlight
                for _ in range(HIGHLIGHT_BLINKS):
                    temp = frame.copy()
                    cv2.rectangle(temp, (x1,y1),(x2,y2), (0,255,255), 4)
                    cv2.putText(temp, "SELECTING (Enter to skip)", (x1, y1-10),
                                cv2.FONT_HERSHEY_SIMPLEX,0.8,(0,255,255),2)
                    cv2.imshow("Semi-Auto Labeling (v10)", temp); cv2.waitKey(180)
                    cv2.imshow("Semi-Auto Labeling (v10)", frame); cv2.waitKey(120)

                pid = input(f"🧍 Enter ID for Center=({det['cx']},{det['cy']}) → ").strip()
                if pid == "":
                    print("⏩ Skipped — no ID assigned.")
                    continue  # unlabeled বাদ

                col = color_for_id(pid)
                cv2.rectangle(frame,(x1,y1),(x2,y2),col,2)
                cv2.putText(frame,f"ID {pid}",(x1+5,y1-8),cv2.FONT_HERSHEY_SIMPLEX,0.7,col,2)
                cv2.imshow("Semi-Auto Labeling (v10)", frame); cv2.waitKey(200)
                manual_entries.append({"id": pid, "x": det["cx"], "y": det["cy"], "box": [x1,y1,x2,y2]})

        # 3) save current frame labels (only labeled persons)
        if manual_entries:
            frame_id_data[str(frame_idx)] = manual_entries
            print("✅ IDs saved for this frame.")

            # accumulate GT rows for CSV
            for e in manual_entries:
                bx = e["box"]
                gt_rows.append({
                    "frame": frame_idx,
                    "person_id": str(e["id"]),
                    "center_x": int(e["x"]),
                    "center_y": int(e["y"]),
                    "box_x1": int(bx[0]),
                    "box_y1": int(bx[1]),
                    "box_x2": int(bx[2]),
                    "box_y2": int(bx[3]),
                })
        else:
            print("⚪ No persons labeled — skipping frame.")

        # optional: save labeled image
        if SAVE_LABELED_FRAMES:
            cv2.imwrite(os.path.join(OUTPUT_DIR,"labeled_frames",f"frame_{frame_idx:06d}.jpg"), frame)

        # auto-save progress
        if frame_idx % SAVE_INTERVAL == 0:
            with open(OUTPUT_JSON, "w") as f:
                json.dump(frame_id_data, f, indent=4)
            # also partial CSV save
            if gt_rows:
                pd.DataFrame(gt_rows).to_csv(OUTPUT_CSV, index=False)
            print(f"💾 Auto-saved at frame {frame_idx}")

        frame_idx += 1
        pbar.update(1)

cap.release()
cv2.destroyAllWindows()
pbar.close()

# =========================
# Final Save (JSON + Word + CSV)
# =========================
with open(OUTPUT_JSON, "w") as f:
    json.dump(frame_id_data, f, indent=4)

# Word report
doc = Document()
doc.add_heading("Manual ID Labeling Summary (v10)", level=1)
table = doc.add_table(rows=1, cols=6)
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Frame","Person ID","Center (x,y)"
hdr[3].text, hdr[4].text, hdr[5].text = "Box x1","Box y1","Box x2,y2"

for fkey in sorted(frame_id_data.keys(), key=lambda x:int(x)):
    for o in frame_id_data[fkey]:
        row = table.add_row().cells
        row[0].text = fkey
        row[1].text = str(o["id"])
        row[2].text = f"({int(o['x'])},{int(o['y'])})"
        row[3].text = str(int(o["box"][0]))
        row[4].text = str(int(o["box"][1]))
        row[5].text = f"{int(o['box'][2])},{int(o['box'][3])}"

doc.save(OUTPUT_DOCX)

# Ground Truth CSV
if gt_rows:
    pd.DataFrame(gt_rows).to_csv(OUTPUT_CSV, index=False)
else:
    pd.DataFrame(columns=["frame","person_id","center_x","center_y","box_x1","box_y1","box_x2","box_y2"]).to_csv(OUTPUT_CSV, index=False)

print(f"\n✅ Done!\n📄 JSON : {OUTPUT_JSON}\n📝 DOCX : {OUTPUT_DOCX}\n📊 CSV  : {OUTPUT_CSV}")
if SAVE_LABELED_FRAMES:
    print(f"🖼️ Labeled frames: {os.path.join(OUTPUT_DIR,'labeled_frames')}")
